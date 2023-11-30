# -*- coding: utf-8 -*-

'''Collection of :py:class:`~pdf2docx.layout.Section` instances.
'''

from docx.enum.section import WD_SECTION
from docx.shared import Pt
from ..common.Collection import BaseCollection
from ..common.docx import reset_paragraph_format
from .Section import Section
from ..common import constants
from ..common.share import rgb_component_from_name


class Sections(BaseCollection):

    def restore(self, raws:list):
        """Restore sections from source dicts."""
        self.reset()
        for raw in raws:
            section = Section().restore(raw)
            self.append(section)
        return self


    def parse(self, **settings):
        '''Parse layout under section level.'''
        for section in self: section.parse(**settings)
        return self


    def make_docx(self, doc):
        '''Create sections in docx.'''
        if not self: return

        # mark paragraph index before creating current page
        n = len(doc.paragraphs)

        def create_dummy_paragraph_for_section(section):
            p = doc.add_paragraph()
            line_height = min(section.before_space, 11)
            pf = reset_paragraph_format(p, line_spacing=Pt(line_height))
            pf.space_after = Pt(section.before_space-line_height)

        # ---------------------------------------------------
        # first section
        # ---------------------------------------------------
        # vertical position: add dummy paragraph only if before space is required
        section = self[0]
        if section.before_space > constants.MINOR_DIST:
            create_dummy_paragraph_for_section(section)

        # create first section
        if section.num_cols==2:
            doc.add_section(WD_SECTION.CONTINUOUS)
        section.make_docx(doc)

        # ---------------------------------------------------
        # more sections
        # ---------------------------------------------------
        for section in self[1:]:
            # create new section symbol
            doc.add_section(WD_SECTION.CONTINUOUS)

            # set after space of last paragraph to define the vertical
            # position of current section
            # NOTE: the after space doesn't work if last paragraph is 
            # image only (without any text). In this case, set after
            # space for the section break.
            p = doc.paragraphs[-2] # -1 is the section break
            if not p.text.strip() and 'graphicData' in p._p.xml:
                p = doc.paragraphs[-1]
            pf = p.paragraph_format
            pf.space_after = Pt(section.before_space)

            # section content
            section.make_docx(doc)

        # ---------------------------------------------------
        # create floating images
        # ---------------------------------------------------
        # lazy: assign all float images to first paragraph of current page
        for image in self.parent.float_images:
            image.make_docx(doc.paragraphs[n])


    def plot(self, page, init_color=None):
        '''Plot all section blocks for debug purpose.'''
        last_block_color = None
        for section in self:
            column_continus = False
            if section.num_cols == 2:
                if section[0].is_agjacent_with(section[1], line_break_free_space_ratio=0.1,
                                               new_paragraph_free_space_ratio=0.85):
                    column_continus = True
            last_color = None # color for the last block of left column
            for i, column in enumerate(section):
                for j, block in enumerate(column.blocks):
                    if block.is_header or block.is_footer:
                        continue
                    if i == 0 and j == 0 and init_color is not None:
                        cur_color = init_color
                    else:
                        cur_color = rgb_component_from_name()
                    if block.is_text_block and i == 1 and j == 0 and column_continus:
                        last_block_color = block.plot(page, color=last_color)
                    elif block.is_text_block:
                        last_block_color = block.plot(page, color=cur_color)
                    else:
                        last_block_color = block.plot(page)
                    last_color = cur_color
                # column.plot(page, stroke=(1,1,0), width=1.5) # column bbox
                # column.blocks.plot(page) # blocks
        return last_block_color
